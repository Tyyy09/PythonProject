# Georgian College
# Coding Lab11 working with Word
# COMP 1112 - F25
# Student Name: Ichty Te
# ID: 200626964

import docx
import openpyxl

# Load Excel and Word files
sheet = openpyxl.load_workbook('D:\\week12\\test.xlsx')
doc = docx.Document("D:\\week12\\test.docx")
spreadsheet = sheet.active

columns = ["A", "B", "C"]

headers = []
paragraphs = []

# Collect headers and paragraphs
for col in columns:
    headerCell = spreadsheet[col + "1"]
    paraCell   = spreadsheet[col + "2"]
    headers.append(headerCell.value)
    paragraphs.append(paraCell.value)

# Find longest and shortest indices
longestIndex = 0
shortestIndex = 0
for i in range(1, len(paragraphs)):
    if len(paragraphs[i]) > len(paragraphs[longestIndex]):
        longestIndex = i
    if len(paragraphs[i]) < len(paragraphs[shortestIndex]):
        shortestIndex = i

# Build new order: longest → middle → shortest
orderedHeaders = [headers[longestIndex]]
orderedParagraphs = [paragraphs[longestIndex]]

for i in range(len(paragraphs)):
    if i != longestIndex and i != shortestIndex:
        orderedHeaders.append(headers[i])
        orderedParagraphs.append(paragraphs[i])

orderedHeaders.append(headers[shortestIndex])
orderedParagraphs.append(paragraphs[shortestIndex])

# ✅ Add to Word in new order
for i in range(len(orderedHeaders)):
    doc.add_heading(orderedHeaders[i], 1)
    print("The paragraph of " + orderedHeaders[i] + " : " + orderedParagraphs[i])
    doc.add_paragraph(orderedParagraphs[i])

# ✅ Print info
print("Longest paragraph is under header '" + orderedHeaders[0] +
      "' with " + str(len(orderedParagraphs[0])) + " characters.")
print("Shortest paragraph is under header '" + orderedHeaders[-1] +
      "' with " + str(len(orderedParagraphs[-1])) + " characters.")

print(" ")

# Add image
doc.add_picture('D:\\Modern Minimalist Graffiti Urban Brand Logo.jpg',
                width=docx.shared.Inches(3),
                height=docx.shared.Inches(3))
print("The image has been imported to Word docx successfully.")

# Save Word document
doc.save("D:\\week12\\test.docx")
